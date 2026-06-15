# rewriter.py
import argparse
import re
import shutil
import sys
import warnings
from pathlib import Path

from docx import Document
from tqdm import tqdm

from context_loader import load_context
from docx_parser import (
    ChapterBucket,
    detect_chapters,
    get_paragraph_text,
    replace_paragraph_text,
    clear_paragraph,
)
from llm_client import LLMClient

KEYWORD_MAP = {
    "nhận dạng biển số xe": "phân tích CV và matching với Job Description",
    "YOLO": "SmartCV",
    "OpenCV": "Spring AI",
    "OCR": "LLM",
    "camera giao thông": "nền tảng tuyển dụng",
}

_SUBHEADING_RE = re.compile(r"^Heading [23]$")


def _is_subheading(para) -> bool:
    return bool(_SUBHEADING_RE.match(para.style.name))


def rewrite_chapter1(doc, chapter: ChapterBucket, llm: LLMClient, kb: str) -> None:
    current_section = chapter.title
    for idx in tqdm(chapter.body_indices, desc=f"  Ch1: {chapter.title[:25]}"):
        para = doc.paragraphs[idx]
        if _is_subheading(para):
            current_section = get_paragraph_text(para)
            continue
        original = get_paragraph_text(para).strip()
        if not original:
            continue
        prompt = (
            f"Ngữ cảnh dự án:\n{kb[:20_000]}\n\n"
            f"Mục hiện tại trong báo cáo: {current_section}\n\n"
            f"Đoạn văn gốc cần viết lại:\n{original}\n\n"
            "Viết lại đoạn văn này cho dự án SmartCV. "
            "Giữ nguyên độ dài tương đương. "
            "Chỉ trả về đoạn văn đã viết lại, không thêm giải thích hay tiêu đề."
        )
        result = llm.rewrite(prompt)
        if result:
            replace_paragraph_text(para, result)


def rewrite_chapter2(doc, chapter: ChapterBucket, llm: LLMClient, kb: str) -> None:
    prompt = (
        f"Ngữ cảnh kỹ thuật của dự án SmartCV:\n{kb}\n\n"
        "Viết nội dung học thuật cho chương 'Cơ Sở Lý Thuyết' của đồ án tốt nghiệp "
        "về dự án SmartCV - nền tảng tuyển dụng thông minh. "
        "Trình bày lý thuyết về: React 19, Spring Boot 3, kiến trúc Microservices, "
        "MongoDB, Redis, RabbitMQ, Elasticsearch, Spring AI với Llama 3, "
        "AWS S3, Docker, CI/CD với GitHub Actions, Golang cho Notification Service. "
        "Viết khoảng 800-1000 chữ, chia đoạn rõ ràng, dùng gạch đầu dòng khi liệt kê."
    )
    result = llm.rewrite(prompt)
    if not result or not chapter.body_indices:
        return
    replace_paragraph_text(doc.paragraphs[chapter.body_indices[0]], result)
    for idx in chapter.body_indices[1:]:
        clear_paragraph(doc.paragraphs[idx])


def rewrite_chapter_heading_only(
    doc, chapter: ChapterBucket, llm: LLMClient, kb: str
) -> None:
    heading_para = doc.paragraphs[chapter.heading_index]
    original = get_paragraph_text(heading_para).strip()
    prompt = (
        f"Ngữ cảnh SmartCV:\n{kb[:10_000]}\n\n"
        f"Tiêu đề gốc của chương: {original}\n\n"
        "Viết lại tiêu đề chương này sao cho phản ánh đúng nội dung dự án SmartCV. "
        "Giữ nguyên số chương (VD: CHƯƠNG 3). "
        "Chỉ trả về tiêu đề, không giải thích."
    )
    result = llm.rewrite(prompt)
    if result:
        replace_paragraph_text(heading_para, result.strip())
    for idx in chapter.body_indices:
        clear_paragraph(doc.paragraphs[idx])


def apply_keyword_replacements(doc) -> None:
    def _replace_in_para(para):
        text = get_paragraph_text(para)
        new_text = text
        for old, new in KEYWORD_MAP.items():
            new_text = new_text.replace(old, new)
        if new_text != text:
            replace_paragraph_text(para, new_text)

    for para in doc.paragraphs:
        _replace_in_para(para)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _replace_in_para(para)


def main():
    parser = argparse.ArgumentParser(
        description="Rewrite a graduation thesis .docx for the SmartCV project"
    )
    parser.add_argument("--context", required=True, help="Path to context dir or .md file")
    parser.add_argument(
        "--template",
        default="docs/BAOCAODOANTOTNGHIEP.docx",
        help="Source .docx template",
    )
    parser.add_argument(
        "--output",
        default="docs/phamlucchuong.docx",
        help="Output .docx path",
    )
    args = parser.parse_args()

    template = Path(args.template)
    output = Path(args.output)

    if not template.exists():
        print(f"Error: template not found: {template}", file=sys.stderr)
        sys.exit(1)

    print("Loading context...")
    try:
        kb = load_context(args.context)
    except (FileNotFoundError, ValueError) as e:
        print(f"Error: {e}", file=sys.stderr)
        sys.exit(1)
    print(f"Context: {len(kb):,} chars loaded")

    output.parent.mkdir(parents=True, exist_ok=True)
    print(f"Cloning {template.name} → {output}")
    shutil.copy(template, output)

    doc = Document(str(output))
    chapters = detect_chapters(doc)
    print(f"Detected {len(chapters)} chapter(s):")
    for c in chapters:
        print(f"  [{c.heading_index}] {c.title[:60]} ({len(c.body_indices)} body paragraphs)")

    llm = LLMClient()
    print(f"LLM mode: {llm.mode}")

    print("Applying keyword replacements...")
    apply_keyword_replacements(doc)

    for chapter in chapters:
        title_upper = chapter.title.upper()
        if re.match(r"CHƯƠNG\s+1\b", title_upper):
            print(f"→ CHƯƠNG 1: keep subheadings, rewrite body paragraphs")
            rewrite_chapter1(doc, chapter, llm, kb)
        elif re.match(r"CHƯƠNG\s+2\b", title_upper):
            print(f"→ CHƯƠNG 2: generate full academic content")
            rewrite_chapter2(doc, chapter, llm, kb)
        elif re.match(r"CHƯƠNG\s+[34]\b", title_upper):
            print(f"→ {chapter.title[:40]}: rewrite heading, clear body")
            rewrite_chapter_heading_only(doc, chapter, llm, kb)

    print(f"Saving → {output}")
    doc.save(str(output))
    print("Done.")


if __name__ == "__main__":
    main()
