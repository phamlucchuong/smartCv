import pytest
from docx import Document
from docx_parser import (
    is_chapter_heading,
    detect_chapters,
    get_paragraph_text,
    replace_paragraph_text,
    clear_paragraph,
    ChapterBucket,
)


def _make_doc_with_styles():
    """Build a minimal Document with chapters for testing."""
    doc = Document()
    # Heading 1 → chapter
    h1 = doc.add_paragraph("CHƯƠNG 1: GIỚI THIỆU")
    h1.style = doc.styles["Heading 1"]
    doc.add_paragraph("Nội dung chương 1.")
    doc.add_paragraph("Đoạn 2 chương 1.")
    # Heading 1 → second chapter
    h2 = doc.add_paragraph("CHƯƠNG 2: CƠ SỞ LÝ THUYẾT")
    h2.style = doc.styles["Heading 1"]
    doc.add_paragraph("Nội dung lý thuyết.")
    return doc


def test_is_chapter_heading_by_style():
    doc = Document()
    p = doc.add_paragraph("Some heading")
    p.style = doc.styles["Heading 1"]
    assert is_chapter_heading(p) is True


def test_is_chapter_heading_by_regex():
    doc = Document()
    p = doc.add_paragraph("CHƯƠNG 3: PHÂN TÍCH THIẾT KẾ")
    assert is_chapter_heading(p) is True


def test_is_not_chapter_heading_for_body():
    doc = Document()
    p = doc.add_paragraph("Đây là đoạn văn thông thường.")
    assert is_chapter_heading(p) is False


def test_detect_chapters_finds_all():
    doc = _make_doc_with_styles()
    chapters = detect_chapters(doc)
    assert len(chapters) == 2
    assert chapters[0].title == "CHƯƠNG 1: GIỚI THIỆU"
    assert chapters[1].title == "CHƯƠNG 2: CƠ SỞ LÝ THUYẾT"


def test_detect_chapters_body_indices():
    doc = _make_doc_with_styles()
    chapters = detect_chapters(doc)
    assert chapters[0].body_indices == [1, 2]
    assert chapters[1].body_indices == [4]


def test_get_paragraph_text_concatenates_runs():
    doc = Document()
    p = doc.add_paragraph("Hello")  # add_paragraph("") produces zero runs
    p.add_run(" World")
    assert get_paragraph_text(p) == "Hello World"


def test_replace_paragraph_text_preserves_run0_format():
    doc = Document()
    p = doc.add_paragraph("original")  # non-empty to guarantee run[0] exists
    p.runs[0].bold = True
    p.add_run(" extra")
    replace_paragraph_text(p, "new text")
    assert get_paragraph_text(p) == "new text"
    assert p.runs[0].bold is True  # formatting preserved
    assert p.runs[1].text == ""    # subsequent run cleared


def test_replace_paragraph_text_single_run():
    doc = Document()
    p = doc.add_paragraph("original")
    replace_paragraph_text(p, "replaced")
    assert get_paragraph_text(p) == "replaced"


def test_replace_paragraph_text_no_runs_is_noop():
    doc = Document()
    p = doc.add_paragraph("")
    for run in p.runs:
        run.text = ""
    replace_paragraph_text(p, "new text")  # should not raise
    assert get_paragraph_text(p) == ""


def test_clear_paragraph_empties_all_text():
    doc = Document()
    p = doc.add_paragraph("some content")
    p.add_run(" more")
    clear_paragraph(p)
    assert get_paragraph_text(p) == ""
